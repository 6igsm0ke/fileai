from docx import Document
import io


def parse_docx(content):
    doc = Document(io.BytesIO(content))
    extracted_text = []
    extracted_tables = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            extracted_text.append(text)

    for table in doc.tables:
        rows = []
        for r in table.rows:
            cells = [cell.text.strip() for cell in r.cells]
            rows.append(cells)
        extracted_tables.append(rows)

    return {
        "extracted_text": "\n".join(extracted_text),
        "extracted_tables": extracted_tables,
    }
